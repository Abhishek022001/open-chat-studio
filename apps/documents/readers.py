import logging

import docx
import pypdf
from pydantic import BaseModel, Field

from apps.documents.patch_docx import patch_docx
from apps.files.models import File

logger = logging.getLogger("ocs.documents")

patch_docx()


class FileReadException(Exception):
    pass


class DocumentPart(BaseModel):
    content: str
    metadata: dict = Field(default_factory=dict)


class Document(BaseModel):
    parts: list[DocumentPart] = Field(default_factory=list)
    """List of parts of the document. Could be pages or chunks of text."""
    metadata: dict = Field(default_factory=dict)
    """Arbitrary metadata associated with the document."""

    @classmethod
    def from_file(cls, file: File):
        reader = get_file_content_reader(file.content_type)
        with file.file.open("rb") as fh:
            return reader(fh).with_metadata(
                {
                    "source_file_id": file.id,
                    "source_file_name": file.name,
                    "source_content_type": file.content_type,
                }
            )

    def with_metadata(self, metadata: dict):
        return Document(parts=self.parts, metadata={**self.metadata, **metadata})

    def get_contents_as_string(self):
        return "".join(part.content for part in self.parts)


def get_file_content_reader(content_type) -> callable:
    if content_type in READERS:
        return READERS[content_type]
    mime_class = content_type.split("/")[0]
    if mime_class in READERS:
        return READERS[mime_class]

    logger.warning(f"No reader found for content type {content_type}. Using default text reader.")
    return read_text


def read_text(file_obj) -> Document:
    try:
        return Document(parts=[DocumentPart(content=file_obj.read().decode())])
    except UnicodeDecodeError as e:
        raise FileReadException("Unable to decode file contents to text") from e


def read_pdf(file_obj) -> Document:
    pages = [
        DocumentPart(content=page.extract_text(), metadata={"page": i})
        for i, page in enumerate(pypdf.PdfReader(file_obj).pages)
    ]
    return Document(parts=pages)


def read_docx(file_obj) -> Document:
    return Document(parts=[DocumentPart(content=paragraph.text) for paragraph in docx.Document(file_obj).paragraphs])


READERS = {
    "application/pdf": read_pdf,
    "text": read_text,
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": read_docx,
    None: read_text,
}
